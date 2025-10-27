# --- utils.py ---

import os
import PyPDF2 # For reading PDFs
from docx import Document # For reading DOCX
import pandas as pd
from dotenv import load_dotenv # For loading API keys from .env file
import sqlite3
import json
import re # Regular expressions for parsing AI output
from datetime import datetime
import google.generativeai as genai # Google Gemini library
import spacy # For text processing (summarization, topic extraction)
from docx.shared import Inches # For DOCX export styling

# --- Load API Keys ---
load_dotenv() # Load variables from .env file in the same directory
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") # Get the Google API key

# --- Configure Google Gemini ---
gemini_model = None # Initialize variable to hold the model client
if GOOGLE_API_KEY:
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        # Basic safety settings to block harmful content
        safety_settings=[
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        ]
        # Initialize the specific Gemini model (use the name verified by list_models)
        gemini_model = genai.GenerativeModel(
            model_name='models/gemini-flash-latest', # Or 'models/gemini-pro-latest' etc.
            safety_settings=safety_settings
            )
        print("✅ Google Gemini Client Initialized successfully.") # Confirmation in terminal
    except Exception as e:
        # Print error if initialization fails (e.g., bad key, network issue)
        print(f"⚠️ Error configuring Google Gemini API: {e}")
        gemini_model = None
else:
    # Inform user if the key is missing
    print("⚠️ Warning: GOOGLE_API_KEY not found in .env file. AI features disabled.")
    gemini_model = None


# --- File Reading Functions ---
def read_docx(file_obj):
    """Reads text content from an uploaded DOCX file object."""
    try:
        doc = Document(file_obj)
        # Extract text from each paragraph, skipping empty ones
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(full_text) # Join paragraphs with double newline for structure
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return "" # Return empty string on failure

def read_pdf(file_obj):
    """Reads text content from an uploaded PDF file object."""
    try:
        reader = PyPDF2.PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            # Add extracted text if the page wasn't blank
            if page_text:
                text += page_text + "\n\n" # Add double newline between pages
        if not text.strip():
             print("⚠️ Warning: PDF text extraction resulted in empty string.")
        return text.strip()
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
        return ""

# --- Text Processing Functions (using spaCy) ---
def summarize_text(text, nlp_model, max_sentences=8, target_length=600): # Increased target length slightly
    """Generates a simple extractive summary using the provided spaCy model."""
    if not text or not nlp_model: return "" # Check for valid input
    try:
        # Dynamically adjust spaCy's max length limit based on input + buffer
        nlp_model.max_length = len(text) + 100
        doc = nlp_model(text)
        sentences = [sent.text.strip() for sent in doc.sents] # Get sentences
        summary = ""
        sentence_count = 0
        current_length = 0
        # Build summary sentence by sentence, respecting limits
        for sent in sentences:
            sent_len = len(sent)
            # Add sentence if within limits
            if sentence_count < max_sentences and current_length + sent_len < target_length:
                 summary += sent + " "
                 current_length += sent_len + 1 # Account for space
                 sentence_count += 1
            # Allow one slightly longer sentence if nearing limit but still under
            elif current_length < target_length * 0.8 and sent_len < (target_length - current_length) * 1.5:
                 summary += sent + " "
                 break # Stop after adding one potentially longer sentence
            elif current_length >= target_length:
                 break # Stop if length limit reached
        # Fallback if no sentences fit or text was short
        if not summary: return text[:target_length] + ("..." if len(text) > target_length else "")
        return summary.strip()
    except ValueError as ve: # Catch specific spaCy length errors
        print(f"⚠️ SpaCy max_length likely exceeded during summary: {ve}")
        return text[:target_length] + ("..." if len(text) > target_length else "")
    except Exception as e:
        print(f"❌ Error during summarization: {e}")
        return text[:target_length] + ("..." if len(text) > target_length else "")


def segment_text_by_topic(full_text, nlp_model):
    """
    Segments the text into topics. Placeholder - returns full document.
    Requires the loaded spaCy model (nlp_model).
    """
    if not nlp_model:
        print("⚠️ SpaCy model not loaded, cannot segment text.")
        return {"Full Document": full_text}
    # TODO: Implement actual topic segmentation using nlp_model (e.g., based on headings, keywords, clustering)
    print("ℹ️ Using placeholder topic segmentation (Full Document).")
    return {"Full Document": full_text} # Current placeholder


def extract_topics(text, nlp_model):
     """Extracts potential topics (named entities) using the provided spaCy model."""
     if not text or not nlp_model: return []
     try:
        nlp_model.max_length = len(text) + 100 # Adjust limit
        doc = nlp_model(text)
        # Define relevant entity types
        allowed_labels = {"ORG", "PERSON", "GPE", "PRODUCT", "EVENT", "WORK_OF_ART", "LAW", "FAC", "LOC", "NORP", "LANGUAGE"}
        # Extract, filter by length and word count
        topics = [ent.text.strip() for ent in doc.ents if ent.label_ in allowed_labels and 3 < len(ent.text.strip()) < 30 and len(ent.text.split()) < 5]
        # Use frequency counting for relevance
        from collections import Counter
        topic_counts = Counter(topics)
        # Return the most common unique topics, up to a limit
        unique_topics = [topic for topic, count in topic_counts.most_common(20)]
        return unique_topics
     except ValueError as ve:
         print(f"⚠️ SpaCy max_length likely exceeded during topic extraction: {ve}")
         return []
     except Exception as e:
        print(f"❌ Error extracting topics: {e}")
        return []


# --- AI Question Generation Function ---

def generate_questions_with_ai(context, num_questions, q_types, topic="", nlp_model=None):
    """
    Calls the Google Gemini API to generate questions based on context and parameters.
    Returns a list of dictionaries, each representing a parsed question.
    """
    # Pre-checks for essential components
    if not gemini_model:
        print("❌ Error: Google Gemini model not initialized or API key missing.")
        return []
    if not nlp_model:
        print("❌ Error: SpaCy model needed for summarization is missing.")
        return []

    # Generate summary to use in the prompt
    summary = summarize_text(context, nlp_model, max_sentences=10, target_length=600) # Slightly longer summary
    if not summary:
        print("⚠️ Warning: Could not generate summary from context. Using raw context (potentially truncated).")
        summary = context[:600] # Use truncated context as fallback

    generated_results = [] # List to hold results for all question types

    # Loop through each requested question type (MCQ, Short Answer, Viva)
    for q_type in q_types:
        # Construct the detailed prompt for the AI
        prompt = f"""CONTEXT:
{summary}

INSTRUCTIONS:
Based *only* on the provided CONTEXT about the topic '{topic}', generate exactly {num_questions} high-quality questions of type '{q_type}'.
For each question:
1.  Assign ONE SINGLE appropriate Bloom's Taxonomy level (Remembering, Understanding, Applying, Analyzing, Evaluating, Creating). Provide brief reasoning (max 1 sentence).
2.  Generate the question text based strictly on the CONTEXT. The question should be clear and stand alone.
3.  If 'MCQ', generate 4 distinct plausible options labeled A), B), C), D). Ensure only one option is correct according to the CONTEXT. Indicate the correct answer using "Correct Answer: [Letter]". Ensure options are relevant to the question.
4.  If 'Short Answer', generate a concise, factual answer based directly on the CONTEXT. Avoid generating just 'N/A' unless truly not answerable from context.
5.  If 'Viva', generate an open-ended question suitable for verbal discussion related to the CONTEXT (provide "N/A" as the answer).

OUTPUT FORMAT (Use exactly this structure for each question, including the ### markers):
### Question 1
Bloom Level: [Level Name]
Reasoning: [Brief justification]
Question Type: {q_type}
Question: [Generated Question Text]
Options: (Only for MCQ)
A) [Option A]
B) [Option B]
C) [Option C]
D) [Option D]
Answer: [Correct Answer Text or Letter for MCQ, or N/A for Viva]

### Question 2
Bloom Level: [Level Name]
Reasoning: [Brief justification]
Question Type: {q_type}
Question: [Generated Question Text]
Options: (Only for MCQ)
A) [Option A]
B) [Option B]
C) [Option C]
D) [Option D]
Answer: [Correct Answer Text or Letter for MCQ, or N/A for Viva]

(Continue pattern for all {num_questions} questions)
"""
        try:
            print(f"ℹ️ Sending prompt to Gemini for {topic} - {q_type}...")
            # Configure generation parameters
            generation_config = genai.types.GenerationConfig(
                max_output_tokens=350 * num_questions, # Generous token limit
                temperature=0.6 # Balance focus and slight variation
            )
            # Call the Gemini API
            response = gemini_model.generate_content(
                prompt,
                generation_config=generation_config
                # Safety settings are applied during model initialization
                )

            generated_text = ""
            # Safely access the response text, checking for potential blocks
            try:
                 # Check candidates first for safety reasons
                 if response.candidates and hasattr(response.candidates[0], 'finish_reason') and response.candidates[0].finish_reason != 'STOP':
                      print(f"⚠️ Warning: Gemini content generation potentially stopped or limited. Reason: {response.candidates[0].finish_reason}")
                      if hasattr(response.candidates[0], 'safety_ratings'): print(f"Safety Ratings: {response.candidates[0].safety_ratings}")
                 # Access text from the first part if available
                 if response.parts: generated_text = response.parts[0].text
                 elif hasattr(response, 'text'): generated_text = response.text # Fallback access
                 else: print("⚠️ Warning: Gemini response structure unexpected or empty.")

            except ValueError: # Handle potential value error if response.text access fails (e.g., block)
                 block_reason = "Unknown"; finish_reason = "Unknown"
                 if hasattr(response, 'prompt_feedback') and response.prompt_feedback.block_reason: block_reason = response.prompt_feedback.block_reason
                 if response.candidates and hasattr(response.candidates[0], 'finish_reason'): finish_reason = response.candidates[0].finish_reason
                 print(f"⚠️ Warning: Content potentially blocked by Gemini for {topic} - {q_type}. Block Reason: {block_reason}, Finish Reason: {finish_reason}")
                 generated_text = ""
            except Exception as resp_err:
                 print(f"❌ Error accessing Gemini response text: {resp_err}")
                 generated_text = ""


            # Process the generated text if it's not empty
            if generated_text and generated_text.strip():
                 print(f"\n--- RAW GEMINI OUTPUT ({topic} - {q_type}) ---")
                 print(generated_text)
                 print("------------------------------------------\n")
                 # Parse the structured text into question dictionaries
                 parsed_questions = parse_structured_output(generated_text, q_type, topic)
                 generated_results.extend(parsed_questions)
            else:
                 print(f"⚠️ Warning: Empty or blocked response from Gemini API for {topic} - {q_type}")

        except Exception as e:
             # Catch potential API call errors (rate limits, config issues etc.)
             print(f"❌ Error calling Google Gemini API for {topic} - {q_type}: {e}")

    return generated_results # Return list of parsed question dictionaries


# --- Parsing Function ---
def parse_structured_output(text, expected_q_type, topic):
    """Parses the AI's structured output (using ### Question markers) into a list of dictionaries."""
    questions = []
    # Split by '### Question ' marker
    blocks = re.split(r'### Question \d+\s*', text)
    if not blocks or len(blocks) <= 1:
         print("⚠️ Warning: Could not split AI output into question blocks using '### Question'. Trying newline split.")
         blocks = text.strip().split('\n\n')
         blocks = [b for b in blocks if b.strip() and ("Bloom Level:" in b or "Question:" in b)] # Basic filter
         if not blocks:
             print("⚠️ Fallback newline split also failed.")
             return []

    for block in blocks:
        # Skip potential empty blocks or the part before the first marker if using regex split
        block = block.strip()
        if not block or block.startswith("### Question"): continue

        # Initialize dictionary
        q_data = {
            "topic": topic, "bloom_level": "Unknown", "reasoning": "",
            "question_type": expected_q_type, # Default to requested type
            "question": "", "options": None, "answer": ""
        }

        try:
            # Extract fields using non-greedy matching (.+?) and flags
            bloom = re.search(r"Bloom Level:\s*(.+?)(?:\n|\Z)", block, re.IGNORECASE)
            if bloom: q_data["bloom_level"] = bloom.group(1).strip()

            reason = re.search(r"Reasoning:\s*(.+?)(?:\n|\Z)", block, re.IGNORECASE)
            if reason: q_data["reasoning"] = reason.group(1).strip()

            q_type = re.search(r"Question Type:\s*(.+?)(?:\n|\Z)", block, re.IGNORECASE)
            if q_type: q_data["question_type"] = q_type.group(1).strip() # Use reported type

            q_match = re.search(r"Question:\s*(.+?)(\n\s*Options:|\n\s*Answer:|\Z)", block, re.IGNORECASE | re.DOTALL)
            if q_match: q_data["question"] = q_match.group(1).strip()

            # --- Specific Parsing based on Question Type ---
            current_q_type = q_data["question_type"] # Use the type reported by AI

            if current_q_type == "MCQ":
                options_dict = {}
                options_block_match = re.search(r"Options:\s*(.*?)(?=\n\s*Answer:|\Z)", block, re.IGNORECASE | re.DOTALL)
                if options_block_match:
                    options_text = options_block_match.group(1).strip()
                    option_matches = re.findall(r"([A-D])\)\s*(.+)", options_text, re.IGNORECASE | re.MULTILINE)
                    for label, opt_text in option_matches:
                         options_dict[label.upper()] = opt_text.strip() # Store as { 'A': 'Text A', ... }
                q_data["options"] = options_dict if options_dict else None

                # Find answer letter (specifically look for "Correct Answer: X")
                ans_match = re.search(r"Correct Answer:\s*([A-D])", block, re.IGNORECASE)
                if ans_match:
                    q_data["answer"] = ans_match.group(1).strip().upper() # Store just the letter
                else: # Fallback to just Answer: X
                    ans_match_fallback = re.search(r"Answer:\s*([A-D])", block, re.IGNORECASE)
                    if ans_match_fallback: q_data["answer"] = ans_match_fallback.group(1).strip().upper()
                    else: q_data["answer"] = "Not Found"

            elif current_q_type == "Short Answer":
                ans_match = re.search(r"Answer:\s*(.+)", block, re.IGNORECASE | re.DOTALL)
                if ans_match:
                     ans_text = ans_match.group(1).strip()
                     # Clean up potential extra markers
                     ans_text = re.sub(r'\n### Question.*', '', ans_text, flags=re.DOTALL).strip()
                     # If answer is just N/A, maybe flag it as not found from context
                     if ans_text.upper() == "N/A":
                         q_data["answer"] = "Not Found in Context"
                         print(f"ℹ️ Short Answer response was 'N/A' for: {q_data['question'][:50]}...")
                     else: q_data["answer"] = ans_text
                else: q_data["answer"] = "Not Found"

            elif current_q_type == "Viva":
                 q_data["answer"] = "N/A" # Explicitly set N/A

            # --- Validation ---
            if q_data["question"] and q_data["question_type"] and q_data["bloom_level"] != "Unknown":
                 if q_data["question_type"] != "Viva" and q_data["answer"] in ["Not Found", "Not Found in Context"]:
                      print(f"⚠️ Warning: Parsed '{q_data['question_type']}' but answer not found/parsed for: {q_data['question'][:50]}...")
                 questions.append(q_data)
            else:
                 print(f"ℹ️ Skipping incomplete parsed block: Missing Q/Type/Bloom.\nBlock Start: {block[:150]}...")

        except Exception as e:
            print(f"❌ Error parsing block: {e}\nBlock content:\n---\n{block}\n---")

    return questions


# --- Question Bank Functions ---
DB_PATH = "database/question_bank.db" # Define path for database file

def init_question_bank_db():
    """Initializes the SQLite question bank database and table."""
    try:
        os.makedirs(os.path.dirname(DB_PATH), exist_ok=True) # Ensure directory exists
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        # Create table schema
        c.execute('''CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT, topic TEXT, question_type TEXT,
            bloom_level TEXT, reasoning TEXT, question TEXT UNIQUE NOT NULL,
            options TEXT, answer TEXT, difficulty TEXT, added_timestamp TEXT NOT NULL
        )''')
        conn.commit()
        conn.close()
    except Exception as e: print(f"❌ Error initializing question bank DB: {e}")

init_question_bank_db() # Initialize DB when script is loaded

def save_to_question_bank(questions_list):
    """Saves a list of question dictionaries to the SQLite DB, handles duplicates."""
    if not questions_list: return 0
    saved_count = 0
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        timestamp = datetime.now().isoformat()
        for q_data in questions_list:
             level = q_data.get('bloom_level', '').lower()
             difficulty = 'Hard' if level in ['creating', 'evaluating'] else ('Medium' if level in ['applying', 'analyzing'] else 'Easy')
             options_json = json.dumps(q_data.get('options')) if q_data.get('options') else None
             question_text = q_data.get('question', '')
             if not question_text: continue
             try:
                c.execute("""INSERT INTO questions
                             (topic, question_type, bloom_level, reasoning, question, options, answer, difficulty, added_timestamp)
                             VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                          (q_data.get('topic', 'N/A'), q_data.get('question_type', 'N/A'), q_data.get('bloom_level', 'N/A'), q_data.get('reasoning', ''),
                           question_text, options_json, str(q_data.get('answer', '')), difficulty, timestamp))
                saved_count += 1
             except sqlite3.IntegrityError: print(f"ℹ️ Skipping duplicate: {question_text[:50]}...")
             except Exception as ie: print(f"❌ Error inserting question: {ie}")
        conn.commit()
        conn.close()
        print(f"ℹ️ Attempted {len(questions_list)}, Added {saved_count} new questions to bank.")
        return saved_count
    except Exception as e:
        print(f"❌ Error saving to question bank: {e}")
        return 0

def load_question_bank():
    """Loads the entire question bank into a pandas DataFrame."""
    try:
        if not os.path.exists(DB_PATH):
             print("ℹ️ Question bank database file not found. Returning empty.")
             return pd.DataFrame()
        conn = sqlite3.connect(DB_PATH)
        query = "SELECT id, topic, question_type, bloom_level, question, answer, difficulty, added_timestamp FROM questions ORDER BY added_timestamp DESC"
        df = pd.read_sql_query(query, conn)
        conn.close()
        if not df.empty and 'added_timestamp' in df.columns:
             try: df['added_timestamp'] = pd.to_datetime(df['added_timestamp']).dt.strftime('%Y-%m-%d %H:%M')
             except Exception as date_err: print(f"⚠️ Warning: Could not format timestamp: {date_err}")
        return df
    except Exception as e:
        print(f"❌ Error loading question bank: {e}")
        return pd.DataFrame()

# --- Export Functions ---

def export_to_csv(questions_list, file_path_base, include_answers=True):
    """
    Exports the list of question dictionaries to a CSV file.
    Generates solved or unsolved versions based on include_answers.
    """
    if not questions_list: return None
    try:
        processed_data = []
        for q_data in questions_list:
             level = q_data.get('bloom_level', '').lower()
             difficulty = 'Hard' if level in ['creating', 'evaluating'] else ('Medium' if level in ['applying', 'analyzing'] else 'Easy')
             row = {'topic': q_data.get('topic', 'N/A'), 'question_type': q_data.get('question_type', 'N/A'),
                    'bloom_level': q_data.get('bloom_level', 'N/A'), 'difficulty': difficulty,
                    'question': q_data.get('question', ''), 'option_a': '', 'option_b': '',
                    'option_c': '', 'option_d': '',
                    'answer': str(q_data.get('answer', '')) if include_answers else '', # Include answer conditionally
                    'reasoning': q_data.get('reasoning', '')}
             if row['question_type'] == 'MCQ' and q_data.get('options') and isinstance(q_data.get('options'), dict):
                  options = q_data.get('options', {})
                  row['option_a'] = options.get('A', '')
                  row['option_b'] = options.get('B', '')
                  row['option_c'] = options.get('C', '')
                  row['option_d'] = options.get('D', '')
             processed_data.append(row)

        df = pd.DataFrame(processed_data)
        # Define columns based on whether answers are included
        export_cols = ['topic', 'question_type', 'bloom_level', 'difficulty', 'question', 'option_a', 'option_b', 'option_c', 'option_d']
        if include_answers: export_cols.append('answer')
        export_cols.append('reasoning') # Keep reasoning optionally

        df_export = df[[col for col in export_cols if col in df.columns]]

        # Construct filename
        suffix = "_solved" if include_answers else "_unsolved"
        file_path = f"{file_path_base}{suffix}.csv"

        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        df_export.to_csv(file_path, index=False, encoding='utf-8-sig')
        print(f"✅ Exported {len(df_export)} questions to {file_path}")
        return file_path
    except Exception as e:
        print(f"❌ Error exporting to CSV: {e}")
        return None

def export_to_docx(questions_list, file_path_base, include_answers=True):
    """
    Exports the list of question dictionaries to a DOCX file.
    Generates solved or unsolved versions based on include_answers.
    """
    if not questions_list: return None
    try:
        doc = Document()
        # --- Document Setup ---
        doc.add_heading('Note2Quiz Generated Questions', level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total Questions: {len(questions_list)}")
        if not include_answers: doc.add_paragraph("Version: Unsolved (No Answer Key)")
        doc.add_page_break()

        # --- Question Section ---
        for i, q_data in enumerate(questions_list, 1):
            if i > 1: doc.add_page_break() # Page break between questions
            doc.add_heading(f"Question {i}", level=2)
            # Add metadata
            doc.add_paragraph(f"Topic: {q_data.get('topic', 'N/A')}")
            doc.add_paragraph(f"Type: {q_data.get('question_type', 'N/A')}")
            level = q_data.get('bloom_level', '').lower()
            difficulty = 'Hard' if level in ['creating', 'evaluating'] else ('Medium' if level in ['applying', 'analyzing'] else 'Easy')
            doc.add_paragraph(f"Bloom's Level: {q_data.get('bloom_level', 'N/A')} ({difficulty})")
            # doc.add_paragraph(f"Reasoning: {q_data.get('reasoning', 'N/A')}") # Optional

            # Add question text (bold label)
            p_q = doc.add_paragraph(); p_q.add_run("Question: ").bold = True; p_q.add_run(str(q_data.get('question', 'N/A')))

            # Add options for MCQ (bulleted, indented)
            if q_data.get('question_type') == 'MCQ' and q_data.get('options') and isinstance(q_data.get('options'), dict):
                options = q_data.get('options', {})
                for label in sorted(options.keys()): # Ensure A, B, C, D order
                     para = doc.add_paragraph(f"{label}) {options[label]}", style='List Bullet')
                     para.paragraph_format.left_indent = Inches(0.5) # Indent options

            # Add space only if we are including the answer key later
            if include_answers: doc.add_paragraph()

        # --- Answer Key Section (Conditional) ---
        if include_answers:
            doc.add_page_break()
            doc.add_heading('Answer Key', level=1)
            for i, q_data in enumerate(questions_list, 1):
                 ans_text = str(q_data.get('answer', 'N/A'))
                 # Enhance answer display for MCQs
                 if q_data.get('question_type') == 'MCQ' and q_data.get('options') and ans_text in q_data['options']:
                      ans_text = f"{ans_text}) {q_data['options'][ans_text]}" # Show letter and text
                 # Add question number and answer
                 doc.add_paragraph(f"{i}. {ans_text}")

        # --- Save ---
        suffix = "_solved" if include_answers else "_unsolved"
        file_path = f"{file_path_base}{suffix}.docx"
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        doc.save(file_path)
        print(f"✅ Exported {len(questions_list)} questions to {file_path}")
        return file_path
    except Exception as e:
        print(f"❌ Error exporting to DOCX: {e}")
        return None

# --- NEW FUNCTION ---
def clear_question_bank():
    """Deletes all records from the questions table in the database."""
    try:
        if not os.path.exists(DB_PATH):
             print("ℹ️ Question bank database file not found. Nothing to clear.")
             return True # Indicate success as there's nothing to do

        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("DELETE FROM questions") # Deletes all rows
        conn.commit()
        # Optional: Vacuum to shrink file size after deletion
        # c.execute("VACUUM")
        # conn.commit()
        conn.close()
        print("✅ Question bank cleared successfully.")
        return True # Indicate success
    except Exception as e:
        print(f"❌ Error clearing question bank: {e}")
        return False # Indicate failure